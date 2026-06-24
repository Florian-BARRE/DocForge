# ====== Code Summary ======
# DocxLayoutHelper — low-level Word layout operations that python-docx does not expose via its
# high-level API (multi-column sections, running headers/footers, landscape sections, nested tables,
# multi-level lists). Centralised here so the builder stays readable and the oxml fiddling lives in
# one place. These exercise the COMPLEX layout paths Gotenberg/Docling must survive.

# ====== Standard Library Imports ======
from __future__ import annotations

from typing import Any

# ====== Third-Party Library Imports ======
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class DocxLayoutHelper:
    """Static helpers for complex Word layout (columns, headers/footers, landscape, nested tables)."""

    def __new__(cls, *args: object, **kwargs: object) -> None:
        raise TypeError("DocxLayoutHelper is a static-only class and cannot be instantiated.")

    @staticmethod
    def set_columns(section: Any, num: int = 2, space_twips: int = 425) -> None:
        """
        Render a section's body in ``num`` newspaper-style columns.

        Args:
            section: A python-docx Section.
            num (int): Number of columns.
            space_twips (int): Gutter width in twips (1/1440 inch).
        """
        # 1. Find or create the <w:cols> element in the section properties
        sect_pr = section._sectPr
        cols = sect_pr.find(qn("w:cols"))
        if cols is None:
            cols = OxmlElement("w:cols")
            sect_pr.append(cols)
        # 2. Set the column count + gutter
        cols.set(qn("w:num"), str(num))
        cols.set(qn("w:space"), str(space_twips))

    @staticmethod
    def add_section(doc: Any, *, landscape: bool = False) -> Any:
        """
        Append a new page-break section and optionally make it landscape.

        Args:
            doc: The python-docx Document.
            landscape (bool): When True, set landscape orientation and swap page dimensions.

        Returns:
            The new Section.
        """
        # 1. Start a fresh section on a new page
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        # 2. Flip to landscape by swapping width/height (python-docx does not swap automatically)
        if landscape:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
        return section

    @staticmethod
    def set_running_header(section: Any, text: str) -> None:
        """Set a non-inherited running header on a section."""
        header = section.header
        header.is_linked_to_previous = False
        header.paragraphs[0].text = text

    @staticmethod
    def set_running_footer(section: Any, text: str) -> None:
        """Set a non-inherited running footer on a section."""
        footer = section.footer
        footer.is_linked_to_previous = False
        footer.paragraphs[0].text = text

    @staticmethod
    def nested_table(cell: Any, headers: list[str], rows: list[list[str]]) -> None:
        """
        Add a small table INSIDE a table cell (a layout torture-test for parsers).

        Args:
            cell: The parent table cell.
            headers (list[str]): Header row of the nested table.
            rows (list[list[str]]): Data rows of the nested table.
        """
        # 1. Create the inner table with a header row
        table = cell.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for col, text in enumerate(headers):
            table.cell(0, col).text = text
        # 2. Append the data rows
        for row in rows:
            cells = table.add_row().cells
            for col, value in enumerate(row):
                cells[col].text = str(value)

    @staticmethod
    def multilevel_list(doc: Any, items: list[str]) -> None:
        """
        Add a 3-level bulleted list (cycles List Bullet / List Bullet 2 / List Bullet 3).

        Args:
            doc: The python-docx Document.
            items (list[str]): The leaf texts; level deepens then resets every three items.
        """
        # 1. Cycle bullet styles to produce a genuinely nested list structure
        styles = ["List Bullet", "List Bullet 2", "List Bullet 3"]
        for i, item in enumerate(items):
            doc.add_paragraph(item, style=styles[i % len(styles)])

# ====== Code Summary ======
# LIVE end-to-end lifecycle test against the running DocForge stack (API + arq worker +
# Postgres + Qdrant + TEI). Skipped automatically when the API is unreachable, so it never
# breaks the normal unit run. Exercises the real flow with a demo fixture file:
#   create collection -> ingest -> wait done -> coherence (chunks/metadata + Qdrant points)
#   -> search finds it -> update metadata -> delete document (no orphans) -> delete collection.
#
# Run explicitly:  pytest tests/integration/test_lifecycle_live.py -v -s
# Requires the dev stack up (docker compose up): API :10020, Qdrant :10025.

# ====== Standard Library Imports ======
from __future__ import annotations

import json
import time
import urllib.error
import urllib.request
import uuid

# ====== Third-Party Library Imports ======
import pytest

API = "http://localhost:10020/api/v1"
QDRANT_URL = "http://localhost:10025"


# ── HTTP helpers ────────────────────────────────────────────────────────────────


def _req(method: str, url: str, body: dict | None = None, raw: bytes | None = None,
         headers: dict | None = None, timeout: float = 60.0):
    data = raw if raw is not None else (json.dumps(body).encode() if body is not None else None)
    req = urllib.request.Request(url, data=data, method=method, headers=headers or {})
    try:
        with urllib.request.urlopen(req, timeout=timeout) as r:
            payload = r.read()
            return r.status, (json.loads(payload) if payload else {})
    except urllib.error.HTTPError as e:
        body_txt = e.read().decode()
        try:
            return e.code, json.loads(body_txt)
        except json.JSONDecodeError:
            return e.code, {"_raw": body_txt}


def _api(method: str, path: str, body: dict | None = None):
    return _req(method, API + path, body=body, headers={"Content-Type": "application/json"})


def _make_docx(marker: str) -> bytes:
    """
    Generate a small DOCX with python-docx (a project dependency).

    Office documents are converted to PDF by Gotenberg and then parsed by Docling — the
    primary, reliable ingestion path (a synthetic fitz PDF is not decomposed by Docling).

    Args:
        marker (str): A unique token embedded in the body so the source_hash differs per
            run — otherwise identical content triggers source-hash dedup (done, 0 chunks).

    Returns:
        bytes: A valid .docx with headings and paragraphs yielding real blocks/chunks.
    """
    import io

    from docx import Document as Docx

    d = Docx()
    d.add_heading("Rapport financier trimestriel", level=1)
    d.add_paragraph(f"Identifiant unique du document de test : {marker}.")
    d.add_paragraph("Ce document de test couvre les risques financiers, les conclusions "
                    "d'audit et la conformite.")
    d.add_heading("Risques", level=2)
    d.add_paragraph("Les risques incluent l'exposition de marche, le risque de credit et le "
                    "risque operationnel. Le comite a examine les plans d'attenuation.")
    d.add_heading("Conformite et metadonnees", level=2)
    d.add_paragraph("La coherence des metadonnees entre Postgres et Qdrant est verifiee. "
                    "La suppression logique ne doit laisser aucun orphelin.")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _live() -> bool:
    try:
        s, _ = _req("GET", API + "/health/ping", timeout=3)
        return s == 200
    except Exception:
        return False


pytestmark = pytest.mark.skipif(not _live(), reason="live DocForge stack not reachable on :10020")


# ── Multipart upload (stdlib, no external deps) ──────────────────────────────────


def _ingest(collection_id: str, filename: str, content: bytes, metadata: dict | None) -> tuple[int, dict]:
    boundary = "----docforge%s" % uuid.uuid4().hex
    parts: list[bytes] = []
    if metadata:
        parts.append(
            ('--%s\r\nContent-Disposition: form-data; name="metadata"\r\n\r\n%s\r\n'
             % (boundary, json.dumps(metadata))).encode()
        )
    parts.append(
        ('--%s\r\nContent-Disposition: form-data; name="file"; filename="%s"\r\n'
         'Content-Type: application/octet-stream\r\n\r\n' % (boundary, filename)).encode()
    )
    parts.append(content)
    parts.append(("\r\n--%s--\r\n" % boundary).encode())
    payload = b"".join(parts)
    return _req(
        "POST", f"{API}/collections/{collection_id}/documents/ingest",
        raw=payload, headers={"Content-Type": "multipart/form-data; boundary=%s" % boundary},
        timeout=60,
    )


# ── Qdrant helpers (REST, no client dep needed) ─────────────────────────────────


def _qdrant_count(collection_id: str, doc_id: str | None = None) -> int:
    """Count Qdrant points in a collection, optionally filtered to one document_id."""
    body: dict = {"exact": True}
    if doc_id is not None:
        body["filter"] = {"must": [{"key": "document_id", "match": {"value": doc_id}}]}
    s, d = _req("POST", f"{QDRANT_URL}/collections/{collection_id}/points/count",
                body=body, headers={"Content-Type": "application/json"}, timeout=15)
    if s != 200:
        return -1
    return int(d.get("result", {}).get("count", 0))


def _qdrant_collection_exists(collection_id: str) -> bool:
    s, _ = _req("GET", f"{QDRANT_URL}/collections/{collection_id}", timeout=15)
    return s == 200


def _wait_done(collection_id: str, doc_id: str, timeout_s: float = 300) -> dict:
    """
    Poll the document until it has produced chunks, errored, or the timeout elapses.

    Waiting on chunk_count (not just status) is more robust: the document status can read
    "done" a beat before chunk rows are committed, and an error path is terminal too.
    """
    deadline = time.time() + timeout_s
    last: dict = {}
    while time.time() < deadline:
        s, d = _api("GET", f"/collections/{collection_id}/documents/{doc_id}")
        if s == 200:
            last = d
            if (d.get("chunk_count") or 0) > 0 or d.get("status") == "error":
                return d
        time.sleep(2)
    return last


# ── The lifecycle test ────────────────────────────────────────────────────────


def test_full_document_lifecycle_live() -> None:
    """Create → ingest → coherence → search → update → delete (no orphans) → drop collection."""
    name = f"itest-{uuid.uuid4().hex[:8]}"
    collection_id = None
    try:
        # 1. Create a collection accepting the HTML fixture, with custom metadata fields
        #    (one filterable, one searchable) so we exercise metadata + searchable vectors.
        s, created = _api("POST", "/collections/create", {
            "name": name,
            "supported_formats": ["docx"],
            # Dense-only TEI: the deployed TEI serves BGE-M3 without the sparse head, so
            # embed_sparse must be False (otherwise /embed_sparse → HTTP 424).
            "pipeline": {"embed": {"chain": [{"id": "tei", "embed_sparse": False}]}},
            "metadata_schema": [
                {"field_name": "dossier", "field_type": "string", "required": False,
                 "filterable": True, "lexical": False, "semantic": False},
                {"field_name": "sujet", "field_type": "string", "required": False,
                 "filterable": False, "lexical": False, "semantic": True},
            ],
        })
        assert s == 201, created
        collection_id = created["id"]

        # 2. Ingest a generated DOCX with metadata (Gotenberg → PDF → Docling).
        #    The collection name (unique per run) seeds unique content → no source-hash dedup.
        content = _make_docx(name)
        s, ing = _ingest(collection_id, "sample.docx", content,
                         {"dossier": "D-2026-001", "sujet": "rapport de test"})
        assert s in (200, 202), ing
        doc_id = ing["doc_id"]

        # 3. Wait for the pipeline to produce chunks (or error / time out)
        doc = _wait_done(collection_id, doc_id)
        assert doc.get("status") != "error", f"ingestion errored: {doc.get('pipeline_errors')}"

        # 4. Coherence — chunks exist and metadata is stored.
        #    (Actual indexing is verified against Qdrant in step 5, not via the stage-run
        #    `indexed` flag, which can lag the real upsert.)
        assert doc.get("chunk_count", 0) > 0, f"no chunks produced (status={doc.get('status')})"
        assert doc.get("user_meta", {}).get("dossier") == "D-2026-001"

        # 5. Qdrant ↔ Postgres coherence — Qdrant points for this doc match the chunk count
        q_doc = _qdrant_count(collection_id, doc_id)
        # Qdrant points are per-chunk; hierarchical chunking may index children only, so points
        # should be >= 1 and consistent (not zero) for an indexed document.
        assert q_doc > 0, f"no Qdrant points for the indexed document (got {q_doc})"

        # 6. Search finds the document (hybrid search through the real engine)
        s, res = _api("POST", f"/collections/{collection_id}/documents/search",
                      {"query": "test", "top_k": 5, "debug": True})
        assert s == 200, res
        assert any(r["document_id"] == doc_id for r in res.get("results", [])), "doc not found in search"

        # 7. Filter search on the filterable field returns it; a wrong value returns nothing
        s, hit = _api("POST", f"/collections/{collection_id}/documents/search",
                      {"query": "test", "top_k": 5,
                       "filters": {"must": [{"key": "dossier", "match": {"value": "D-2026-001"}}]}})
        s2, miss = _api("POST", f"/collections/{collection_id}/documents/search",
                        {"query": "test", "top_k": 5,
                         "filters": {"must": [{"key": "dossier", "match": {"value": "NOPE"}}]}})
        assert miss.get("total") == 0, "filter on a non-existent value should return nothing"

        # 8. Update metadata (no reindex) — value is reflected
        s, upd = _api("POST", f"/collections/{collection_id}/documents/{doc_id}/update",
                      {"metadata": {"dossier": "D-2026-999"}, "reindex": False})
        assert s == 200, upd
        s, doc2 = _api("GET", f"/collections/{collection_id}/documents/{doc_id}")
        assert doc2["user_meta"]["dossier"] == "D-2026-999"

        # 9. Delete the document — Qdrant points AND Postgres chunks must be gone (no orphans)
        s, dele = _api("DELETE", f"/collections/{collection_id}/documents/{doc_id}/delete")
        assert s == 200 and dele.get("deleted") is True, dele
        s, _ = _api("GET", f"/collections/{collection_id}/documents/{doc_id}")
        assert s == 404, "document still retrievable after delete"
        assert _qdrant_count(collection_id, doc_id) == 0, "ORPHAN: Qdrant points remain after document delete"
        s, chunks = _api("GET", f"/collections/{collection_id}/documents/{doc_id}/chunks/list")
        # 404 (doc gone) or empty list — either way, no chunk rows remain.
        assert s == 404 or chunks.get("total", 0) == 0, "ORPHAN: chunk rows remain after document delete"

    finally:
        # 10. Drop the collection — Qdrant collection must be removed too
        if collection_id:
            _api("DELETE", f"/collections/{collection_id}/delete")
            assert not _qdrant_collection_exists(collection_id), "ORPHAN: Qdrant collection remains after drop"

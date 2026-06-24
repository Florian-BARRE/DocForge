# ====== Code Summary ======
# DocxCorpusBuilder — builds a LONG, complex .docx from a multilingual content pack. Composes:
# title + subtitle, running header/footer, abstract, a multi-column section, a long body (contract
# articles or report sections with deep prose), embedded figures with captions, a LANDSCAPE section
# carrying a wide table with a NESTED table, footnote-style notes and a multi-level list. The volume
# of prose stresses chunking; the layout features stress Gotenberg/Docling; the language drives
# detection. Driven by spec.doc_type ("contract"/"report") + spec.language ("fr"/"en"/"es").

# ====== Standard Library Imports ======
from __future__ import annotations

import io

# ====== Third-Party Library Imports ======
from docx import Document
from docx.shared import Inches, Pt, RGBColor

# ====== Local Project Imports ======
from .base import BaseDocumentBuilder
from .content import ContentPack, get_content
from .docx_layout import DocxLayoutHelper
from .image_factory import ImageFactory

# How many body sections/articles to emit (cycling the content pool) — drives length + chunking.
_BODY_BLOCKS = 12


class DocxCorpusBuilder(BaseDocumentBuilder):
    """Builds a long, structurally complex Word document for one (doc_type, language)."""

    def build(self) -> bytes:
        """
        Assemble the .docx and return its bytes.

        Returns:
            bytes: A long, complex .docx document.
        """
        # 1. Resolve the content pack for this document's type + language
        content = get_content(self.spec.doc_type, self.spec.language)
        doc = Document()

        # 2. Title block + running header/footer on the opening section
        self.__add_title(doc, content)
        self.__add_running_chrome(doc, content)

        # 3. Multi-column abstract / governance section (newspaper layout)
        self.__add_columns_section(doc, content)

        # 4. Long single-column body (contract articles or report sections)
        DocxLayoutHelper.add_section(doc)  # back to one column for the body
        if self.spec.doc_type == "contract":
            self.__add_articles(doc, content)
        else:
            self.__add_report_sections(doc, content)

        # 5. Embedded figures with captions
        self.__add_figures(doc, content)

        # 6. Landscape section: wide table with a nested table inside
        self.__add_landscape_table(doc, content)

        # 7. Notes (footnote-style) + a multi-level list
        self.__add_notes_and_list(doc, content)

        # 8. Serialize
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ─── Title + chrome ───────────────────────────────────────────────────────────

    def __add_title(self, doc: Document, content: ContentPack) -> None:
        """Add the title (H0), an italic subtitle and the distinctive abstract paragraph."""
        doc.add_heading(content.title, level=0)
        subtitle = doc.add_paragraph()
        run = subtitle.add_run(content.subtitle)
        run.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        doc.add_paragraph(content.abstract)

    @staticmethod
    def __add_running_chrome(doc: Document, content: ContentPack) -> None:
        """Set a running header (title) and footer (subtitle) on the first section."""
        section = doc.sections[0]
        DocxLayoutHelper.set_running_header(section, content.title)
        DocxLayoutHelper.set_running_footer(section, content.subtitle)

    @staticmethod
    def __add_columns_section(doc: Document, content: ContentPack) -> None:
        """Add a two-column section carrying the (repeated) governance blurb."""
        section = DocxLayoutHelper.add_section(doc)
        DocxLayoutHelper.set_columns(section, num=2)
        for _ in range(3):  # repeat to fill both columns with flowing text
            doc.add_paragraph(content.column_blurb)

    # ─── Body: contract ──────────────────────────────────────────────────────────

    def __add_articles(self, doc: Document, content: ContentPack) -> None:
        """Emit numbered contract articles: heading + clause prose + an occasional sub-list."""
        for i in range(_BODY_BLOCKS):
            doc.add_heading(f"Article {i + 1} — {content.section_title(i)}", level=1)
            doc.add_paragraph(content.clause(i))
            doc.add_paragraph(content.para(i))
            if i % 3 == 2:  # periodic nested list keeps the structure deep
                DocxLayoutHelper.multilevel_list(doc, content.list_items)
            if i % 4 == 3:
                doc.add_page_break()

    # ─── Body: report ────────────────────────────────────────────────────────────

    def __add_report_sections(self, doc: Document, content: ContentPack) -> None:
        """Emit report sections: H1 + prose, a H2 sub-section + prose, periodic page breaks."""
        for i in range(_BODY_BLOCKS):
            doc.add_heading(f"{i + 1}. {content.section_title(i)}", level=1)
            doc.add_paragraph(content.para(i))
            doc.add_paragraph(content.para(i + 1))
            doc.add_heading(content.section_title(i + 1), level=2)
            doc.add_paragraph(content.para(i + 2))
            if i % 4 == 3:
                doc.add_page_break()

    # ─── Figures ─────────────────────────────────────────────────────────────────

    @staticmethod
    def __add_figures(doc: Document, content: ContentPack) -> None:
        """Embed three distinct figures, each with an italicised caption."""
        doc.add_page_break()
        figures = [
            (ImageFactory.bar_chart("KPI"), f"{content.table_caption} — figure"),
            (ImageFactory.diagram(), content.section_title(0)),
            (ImageFactory.photo(), content.section_title(1)),
        ]
        for idx, (png, caption) in enumerate(figures, start=1):
            doc.add_picture(io.BytesIO(png), width=Inches(4.5))
            para = doc.add_paragraph()
            run = para.add_run(f"Figure {idx}. {caption}")
            run.italic = True
            run.font.size = Pt(9)

    # ─── Landscape wide table (with a nested table) ──────────────────────────────

    @staticmethod
    def __add_landscape_table(doc: Document, content: ContentPack) -> None:
        """Add a landscape section with the wide table; nest a small table in the last cell."""
        DocxLayoutHelper.add_section(doc, landscape=True)
        doc.add_heading(content.table_caption, level=2)
        table = doc.add_table(rows=1, cols=len(content.table_headers))
        table.style = "Table Grid"
        for col, head in enumerate(content.table_headers):
            table.cell(0, col).text = head
        for row in content.table_rows:
            cells = table.add_row().cells
            for col, value in enumerate(row):
                cells[col].text = str(value)
        # Nest a 2-row table inside the first data row's last cell — a parser torture-test.
        DocxLayoutHelper.nested_table(
            table.rows[1].cells[-1],
            headers=["Sous-poste", "Valeur"],
            rows=[["A", "1"], ["B", "2"]],
        )

    # ─── Notes + multi-level list ────────────────────────────────────────────────

    @staticmethod
    def __add_notes_and_list(doc: Document, content: ContentPack) -> None:
        """Add a footnote-style notes block (numbered) and a final multi-level bulleted list."""
        DocxLayoutHelper.add_section(doc)
        doc.add_heading(content.section_title(2), level=1)
        for note in content.notes:
            doc.add_paragraph(note, style="List Number")
        DocxLayoutHelper.multilevel_list(doc, content.list_items)
